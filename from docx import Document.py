from docx import Document

doc = Document()
doc.add_heading('CONTRATO DE PRESTAÇÃO DE SERVIÇOS AUTÔNOMOS - CUIDADOR DE IDOSOS', level=1)

texto = """
Pelo presente instrumento particular de contrato, de um lado, o idoso, xxxxxxxxxxxxxxx, inscrito no CPF sob o nº xxxxxxxxxxxxxxxxxx, residente e domiciliado na ............, brasileira, divorciada, do lar, REPRESENTADO por seu curador xxxxxxxxxxxxxxx, brasileiro, solteiro, inscrito no CPF xxxxxxxxxxxxxx, residente e domiciliado na xxxxxxxxxxxxxxxxx, doravante denominado CONTRATANTE e de outro lado xxxxxxxxxxxxxxxxxxxx, brasileira, autônoma, solteira, portador(a) da cédula de identidade R.G. xxxxxxxxxxxxx...

CLÁUSULA PRIMEIRA: DO OBJETO CONTRATUAL
Este contrato destina-se a regular a prestação de serviços de cuidador de idoso de forma autônoma e por tempo INDETERMINADO.

CLÁUSULA SEGUNDA: DA FUNÇÃO
A CONTRATADA se compromete a desempenhar na residência do CONTRATANTE as obrigações inerentes ao cargo de cuidadora compatíveis com sua função, quais sejam:
I - Cuidar, auxiliar e zelar pelo bem-estar, saúde, alimentação e higiene pessoal do Contratante;
II - Comunicar aos familiares imediatamente qualquer problema com o idoso;
III - Ministrar medicamentos prescritos;
IV - Acompanhar o idoso em atividades externas necessárias.

§1º As funções devem ser realizadas com EPIs: luvas, máscaras, álcool em gel etc.
§2º É vedada qualquer intervenção fora da competência técnica da CONTRATADA.

CLÁUSULA TERCEIRA: DO LOCAL DA PRESTAÇÃO E HORÁRIOS
A CONTRATADA prestará serviços na residência da CONTRATANTE na ....................

§2º Após o contrato, a CONTRATADA deverá manter sigilo absoluto sobre a vida pessoal do CONTRATANTE.

CLÁUSULA SÉTIMA: DA RESCISÃO
Pode ser rescindido com aviso prévio de 30 dias. Motivos:
I - Descumprimento de obrigações;
II - Maus tratos ao idoso;
III - Condenação criminal;
IV - Morte do CONTRATANTE ou CONTRATADA.

CLÁUSULA OITAVA: DO VÍNCULO
Não há vínculo empregatício.

CLÁUSULA NONA: DO FORO
Foro eleito: Florianópolis/SC.

Florianópolis, 03 de março de 2021.

Contratante: _________________________
Contratada: _________________________

TESTEMUNHAS:
1ª) ASS: _________________________  NOME: __________________ RG: _______________
2ª) ASS: _________________________  NOME: __________________ RG: _______________
"""

for par in texto.strip().split('\n\n'):
    doc.add_paragraph(par.strip())

doc.save("Contrato_Cuidador_Idosos.docx")
print("Contrato salvo como 'Contrato_Cuidador_Idosos.docx'")
